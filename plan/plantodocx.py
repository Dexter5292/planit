from docx import Document

class plan:
	"Constructor class for the lesson plan."

	def __init__(self):
		self.user = ''
		self.subject = ''
		self.date = ''
		self.grade = ''
		self.time = ''
		self.units = []
		self.topics = []
		self.content = []
		self.activities = []
		self.support = []
		self.recources = []
		self.space = []
		self.homework = []
		self.homeworkDate = []
		self.assessment = []

	def set_user(self, user):
		self.user =user


	def set_subject(self, subject):
		self.subject = subject


	def set_date(self, date):
		self.date = date


	def set_grade(self, grade):
		self.grade = grade


	def set_time(self, time):
		self.time = time


	def set_units(self, units):
		self.units = units


	def set_topics(self, topics):
		""" List parameter with topics from 
		the relevant units {based on user selection} """
		self.topics = topics


	def set_content(self, content):
		""" List parameter with content from 
		the relevant topics {based on user selection} """
		self.content = content

"""
	def generate(self):
		lessonplan = Document()
		sections = lessonplan.sections
		for section in sections:
			section.top_margin = 114300
			section.bottom_margin = 114300
			section.left_margin = 460000
			section.right_margin = 460000
		lessonplan.add_heading("Lesson Plan: %s" % date, level=0)
		table = lessonplan.add_table(rows=4, cols=2)
		lbl_subject = table.cell(0, 0)
		lbl_subject.text = "Subject"
		lbl_date = table.cell(1, 0)
		lbl_date.text = "Date"
		lbl_grade = table.cell(2, 0)
		lbl_grade.text = "Grade/Level"
		lbl_time = table.cell(3, 0)
		lbl_time.text = "Suggested Lesson Time"
		sub_cell = table.cell(0, 1)
		sub_cell.text = "%s" % subject
		date_cell = table.cell(1, 1)
		date_cell.text = date
		grade_cell = table.cell(2, 1)
		grade_cell.text = "{}".format(grade)

		period_time = school_info.objects.get(username=user)
		period_time = period_time.sch_period_length

		time_cell = table.cell(3, 1)
		time_cell.text = "{} minutes".format(period_time)

		lessonplan.save("./%s.docx" % user) """